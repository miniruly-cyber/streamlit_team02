# app.py
# =========================================================
# AI ìê¸°ì†Œê°œì„œ ì½”ì¹­ - ì¹´ì¹´ì˜¤í†¡ ìŠ¤íƒ€ì¼ UI
# =========================================================
# ì„¤ì¹˜: pip install streamlit python-docx reportlab langchain langchain-openai python-dotenv
# ì‹¤í–‰: streamlit run app.py
# =========================================================

import os, io, datetime, json
from typing import Optional, List, Dict
import streamlit as st

# ===== ë¬¸ì„œ ìƒì„± ë¼ì´ë¸ŒëŸ¬ë¦¬ (ì„ íƒ) =====
try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    DOC_LIBS_AVAILABLE = True
except:
    DOC_LIBS_AVAILABLE = False

# ===== LangChain (ì„ íƒ) =====
try:
    from langchain_openai import ChatOpenAI
    from langchain.prompts import ChatPromptTemplate
    from langchain.chains import LLMChain
    LANGCHAIN_AVAILABLE = True
except:
    LANGCHAIN_AVAILABLE = False

# ================= í˜ì´ì§€ ì„¤ì • =================
st.set_page_config(
    page_title="AI ìê¸°ì†Œê°œì„œ ì½”ì¹­",
    page_icon="ğŸ’¬",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ================= ì¹´ì¹´ì˜¤í†¡ ìŠ¤íƒ€ì¼ CSS =================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;600;700&display=swap');
    
    /* ì „ì²´ ë°°ê²½ ë° ê¸°ë³¸ ìŠ¤íƒ€ì¼ */
    .stApp {
        background: #b2c7d9;
        font-family: 'Noto Sans KR', sans-serif;
    }
    
    /* ë©”ì¸ ì»¨í…Œì´ë„ˆ */
    .main .block-container {
        padding: 0;
        max-width: 100%;
        margin: 0;
    }
    
    /* ê¸°ë³¸ íƒ­ ìˆ¨ê¹€ */
    .stTabs [data-baseweb="tab-list"] {
        display: none;
    }
    
    .stTabs [data-baseweb="tab-panel"] {
        padding: 0;
    }
    
    /* ìƒë‹¨ í—¤ë” */
    .chat-header {
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        height: 60px;
        background: rgba(0, 0, 0, 0.85);
        color: white;
        display: flex;
        align-items: center;
        justify-content: center;
        z-index: 1000;
        backdrop-filter: blur(10px);
        border-bottom: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    .header-title {
        font-size: 18px;
        font-weight: 600;
        color: white;
    }
    
    /* ì±„íŒ… ì˜ì—­ */
    .chat-container {
        margin-top: 60px;
        margin-bottom: 120px;
        padding: 20px;
        min-height: calc(100vh - 180px);
        overflow-y: auto;
    }
    
    /* ë©”ì‹œì§€ ë²„ë¸” */
    .msg-row {
        display: flex;
        margin-bottom: 15px;
        align-items: flex-end;
    }
    
    .msg-row.user {
        justify-content: flex-end;
    }
    
    .msg-row.ai {
        justify-content: flex-start;
    }
    
    .msg-bubble {
        max-width: 70%;
        padding: 10px 14px;
        border-radius: 18px;
        font-size: 14px;
        line-height: 1.5;
        word-break: break-word;
        box-shadow: 0 1px 2px rgba(0, 0, 0, 0.15);
        position: relative;
    }
    
    .msg-bubble.user {
        background: #ffeb33;
        color: #000;
        border-top-right-radius: 4px;
    }
    
    .msg-bubble.ai {
        background: white;
        color: #000;
        border-top-left-radius: 4px;
    }
    
    .msg-time {
        font-size: 11px;
        color: #888;
        margin: 0 8px;
        white-space: nowrap;
    }
    
    /* ì„¤ì • í˜ì´ì§€ */
    .settings-container {
        margin-top: 60px;
        margin-bottom: 60px;
        padding: 20px;
        background: white;
        min-height: calc(100vh - 120px);
    }
    
    .settings-section {
        background: #f8f9fa;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 15px;
        border: 1px solid #e9ecef;
    }
    
    .settings-title {
        font-size: 16px;
        font-weight: 600;
        margin-bottom: 15px;
        color: #333;
    }
    
    /* ì €ì¥ì†Œ í˜ì´ì§€ */
    .storage-container {
        margin-top: 60px;
        margin-bottom: 60px;
        padding: 20px;
        background: white;
        min-height: calc(100vh - 120px);
    }
    
    .file-item {
        background: #f8f9fa;
        border-radius: 10px;
        padding: 15px;
        margin-bottom: 10px;
        display: flex;
        justify-content: space-between;
        align-items: center;
        border: 1px solid #e9ecef;
    }
    
    .file-info {
        flex: 1;
    }
    
    .file-name {
        font-weight: 500;
        margin-bottom: 5px;
    }
    
    .file-date {
        font-size: 12px;
        color: #888;
    }
    
    /* ë²„íŠ¼ ìŠ¤íƒ€ì¼ */
    .stButton > button {
        background: #ffeb33;
        color: #000;
        border: none;
        border-radius: 20px;
        padding: 8px 20px;
        font-weight: 500;
        transition: all 0.2s;
        width: 100%;
    }
    
    .stButton > button:hover {
        background: #ffd900;
        color: #000;
    }
    
    /* ì…ë ¥ì°½ ìŠ¤íƒ€ì¼ */
    .stTextInput > div > div > input {
        background: #f5f5f5;
        border: 1px solid #e0e0e0;
        border-radius: 20px;
        padding: 10px 15px;
        font-size: 14px;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #ffeb33 !important;
        box-shadow: 0 0 0 2px rgba(255, 235, 51, 0.2) !important;
    }
    
    /* íŒŒì¼ ì—…ë¡œë“œ ìŠ¤íƒ€ì¼ */
    .stFileUploader > label {
        background: #f8f9fa;
        border: 2px dashed #dee2e6;
        border-radius: 10px;
        padding: 20px;
        text-align: center;
    }
    
    /* ìŠ¤í¬ë¡¤ë°” */
    ::-webkit-scrollbar {
        width: 6px;
    }
    
    ::-webkit-scrollbar-track {
        background: #f1f1f1;
    }
    
    ::-webkit-scrollbar-thumb {
        background: #888;
        border-radius: 3px;
    }
    
    /* selectbox ìŠ¤íƒ€ì¼ */
    .stSelectbox > div > div {
        background: #f5f5f5;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
    }
    
    /* ìŠ¬ë¼ì´ë” ìŠ¤íƒ€ì¼ */
    .stSlider > div > div > div {
        color: #ffeb33;
    }
    
    /* ì •ë³´ ë°•ìŠ¤ ìŠ¤íƒ€ì¼ */
    .stInfo {
        background: #e3f2fd;
        border: 1px solid #bbdefb;
        border-radius: 8px;
    }
    
    /* ì„±ê³µ ë©”ì‹œì§€ ìŠ¤íƒ€ì¼ */
    .stSuccess {
        background: #e8f5e8;
        border: 1px solid #c8e6c9;
        border-radius: 8px;
    }
    
    /* ê²½ê³  ë©”ì‹œì§€ ìŠ¤íƒ€ì¼ */
    .stWarning {
        background: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 8px;
    }
    
    /* ì±„íŒ… ì…ë ¥ ì˜ì—­ ê³ ì • */
    .chat-input-section {
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        background: white;
        padding: 15px;
        border-top: 1px solid #e0e0e0;
        z-index: 999;
    }
    
    /* ë¹ ë¥¸ ë‹µë³€ ë²„íŠ¼ */
    .quick-replies {
        margin-bottom: 10px;
    }
    
    .quick-reply-btn {
        display: inline-block;
        padding: 6px 12px;
        margin: 2px;
        background: white;
        border: 1px solid #e0e0e0;
        border-radius: 15px;
        font-size: 12px;
        cursor: pointer;
        transition: all 0.2s;
    }
    
    .quick-reply-btn:hover {
        background: #ffeb33;
        border-color: #ffeb33;
    }
    
    /* ë°˜ì‘í˜• ë””ìì¸ */
    @media (max-width: 768px) {
        .msg-bubble {
            max-width: 85%;
        }
        
        .settings-container, .storage-container {
            padding: 10px;
        }
        
        .chat-container {
            padding: 10px;
        }
    }
</style>
""", unsafe_allow_html=True)

# ================= ì„¸ì…˜ ì´ˆê¸°í™” =================
if "messages" not in st.session_state:
    st.session_state.messages = []
    st.session_state.messages.append({
        "role": "ai",
        "content": "ì•ˆë…•í•˜ì„¸ìš”! AI ìê¸°ì†Œê°œì„œ ì½”ì¹˜ì…ë‹ˆë‹¤. ë¬´ì—‡ì„ ë„ì™€ë“œë¦´ê¹Œìš”?",
        "time": datetime.datetime.now().strftime("%H:%M")
    })

if "current_tab" not in st.session_state:
    st.session_state.current_tab = "ëŒ€í™”"

if "api_key" not in st.session_state:
    st.session_state.api_key = os.getenv("OPENAI_API_KEY", "")

if "saved_files" not in st.session_state:
    st.session_state.saved_files = []

if "save_format" not in st.session_state:
    st.session_state.save_format = "txt"

if "model_settings" not in st.session_state:
    st.session_state.model_settings = {
        "temperature": 0.7,
        "max_length": 1000,
        "tone": "professional"
    }

# ================= ê°€ì´ë“œë¼ì¸ ì‘ë‹µ =================
def get_guideline_response():
    return """ğŸ“ **AI ìê¸°ì†Œê°œì„œ ì…ë ¥ ê°€ì´ë“œë¼ì¸**

**1. êµ¬ì²´ì ìœ¼ë¡œ ì§ˆë¬¸í•˜ê¸°**
âœ… "ë§ˆì¼€íŒ… ì§ë¬´ ì‹ ì… ìê¸°ì†Œê°œì„œ ë„ì…ë¶€ ì‘ì„±í•´ì¤˜"
âŒ "ìì†Œì„œ ì¨ì¤˜"

**2. ë°°ê²½ ì •ë³´ ì œê³µí•˜ê¸°**
â€¢ ì§€ì› íšŒì‚¬ì™€ ì§ë¬´
â€¢ ë³¸ì¸ì˜ ì£¼ìš” ê²½í—˜
â€¢ ê°•ì¡°í•˜ê³  ì‹¶ì€ ì—­ëŸ‰

**3. íš¨ê³¼ì ì¸ ì§ˆë¬¸ ì˜ˆì‹œ**
â€¢ "ê³ ê° ì„œë¹„ìŠ¤ ê²½í—˜ì„ ì˜ì—…ì§ë¬´ì— ì—°ê²°í•˜ëŠ” ë°©ë²•"
â€¢ "í”„ë¡œì íŠ¸ ê²½í—˜ì„ STAR ê¸°ë²•ìœ¼ë¡œ ì •ë¦¬í•´ì¤˜"
â€¢ "IT ê¸°ì—… ì§€ì›ë™ê¸° ì‘ì„± ë„ì™€ì¤˜"

**4. ì²¨ì‚­ ìš”ì²­ ë°©ë²•**
â€¢ ì‘ì„±í•œ ë‚´ìš© ë³µì‚¬ í›„ "ì´ ë‚´ìš© ì²¨ì‚­í•´ì¤˜"
â€¢ íŒŒì¼ ì—…ë¡œë“œ í›„ "êµ¬ì²´ì„± ë†’ì—¬ì¤˜"
â€¢ "ì´ ë¬¸ì¥ ë” ì„íŒ©íŠ¸ ìˆê²Œ ìˆ˜ì •í•´ì¤˜"

**5. ë‹¨ê³„ë³„ ì ‘ê·¼**
1ï¸âƒ£ ì „ì²´ êµ¬ì¡° ì¡ê¸°
2ï¸âƒ£ ê° ë¬¸ë‹¨ ì‘ì„±
3ï¸âƒ£ í‘œí˜„ ë‹¤ë“¬ê¸°
4ï¸âƒ£ ìµœì¢… ê²€í† 

ğŸ’¡ **Tip**: í•œ ë²ˆì— ëª¨ë“  ê±¸ í•´ê²°í•˜ë ¤ í•˜ì§€ ë§ê³ , ë‹¨ê³„ë³„ë¡œ ì§ˆë¬¸í•˜ì„¸ìš”!"""

# ================= AI ì‘ë‹µ ìƒì„± =================
def get_ai_response(user_input: str, uploaded_file=None) -> str:
    # ê°€ì´ë“œë¼ì¸ ìš”ì²­ ì²´í¬
    guideline_keywords = ["ê°€ì´ë“œ", "ê°€ì´ë“œë¼ì¸", "ë„ì›€ë§", "ì‚¬ìš©ë²•", "ì–´ë–»ê²Œ"]
    if any(keyword in user_input for keyword in guideline_keywords):
        return get_guideline_response()
    
    # í…œí”Œë¦¿ ì‘ë‹µ (API í‚¤ ì—†ì„ ë•Œ)
    if not st.session_state.api_key or not LANGCHAIN_AVAILABLE:
        templates = {
            "default": """ìê¸°ì†Œê°œì„œ ì‘ì„±ì„ ë„ì™€ë“œë¦¬ê² ìŠµë‹ˆë‹¤!

êµ¬ì²´ì ìœ¼ë¡œ ì•Œë ¤ì£¼ì‹œë©´ ë” ì •í™•í•œ ë„ì›€ì„ ë“œë¦´ ìˆ˜ ìˆì–´ìš”:
â€¢ ì–´ë–¤ ì§ë¬´ì— ì§€ì›í•˜ì‹œë‚˜ìš”?
â€¢ ì–´ë–¤ ë¶€ë¶„ì´ ì–´ë ¤ìš°ì‹ ê°€ìš”?
â€¢ íŠ¹ë³„íˆ ê°•ì¡°í•˜ê³  ì‹¶ì€ ê²½í—˜ì´ ìˆë‚˜ìš”?""",
            
            "ì²¨ì‚­": """ìê¸°ì†Œê°œì„œ ì²¨ì‚­ í¬ì¸íŠ¸ë¥¼ ì•Œë ¤ë“œë¦´ê²Œìš”:

âœ… êµ¬ì²´ì ì¸ ìˆ«ìì™€ ì„±ê³¼ í¬í•¨
âœ… ì§ë¬´ì™€ ì—°ê´€ëœ ê²½í—˜ ê°•ì¡°
âœ… ë¬¸ì¥ì€ ê°„ê²°í•˜ê³  ëª…í™•í•˜ê²Œ
âœ… ì§„ì •ì„± ìˆëŠ” ì§€ì›ë™ê¸°

íŒŒì¼ì„ ì—…ë¡œë“œí•˜ê±°ë‚˜ ë‚´ìš©ì„ ë³´ë‚´ì£¼ì‹œë©´ ë” ìì„¸íˆ ë´ë“œë¦´ê²Œìš”!""",
            
            "ì‹œì‘": """ìê¸°ì†Œê°œì„œ ì‘ì„±ì„ ì‹œì‘í•´ë³¼ê¹Œìš”?

**Step 1. ê¸°ë³¸ ì •ë³´**
â€¢ ì§€ì› íšŒì‚¬:
â€¢ ì§€ì› ì§ë¬´:
â€¢ ê²½ë ¥ êµ¬ë¶„: (ì‹ ì…/ê²½ë ¥)

ì´ ì •ë³´ë¥¼ ì•Œë ¤ì£¼ì‹œë©´ ë§ì¶¤í˜•ìœ¼ë¡œ ë„ì™€ë“œë¦´ê²Œìš”!"""
        }
        
        if "ì²¨ì‚­" in user_input or "ìˆ˜ì •" in user_input:
            return templates["ì²¨ì‚­"]
        elif "ì‹œì‘" in user_input or "ì²˜ìŒ" in user_input:
            return templates["ì‹œì‘"]
        else:
            return templates["default"]
    
    # LangChainì„ ì´ìš©í•œ AI ì‘ë‹µ
    try:
        llm = ChatOpenAI(
            api_key=st.session_state.api_key,
            model="gpt-4o-mini",
            temperature=st.session_state.model_settings["temperature"]
        )
        
        system_prompt = f"""ë‹¹ì‹ ì€ ì „ë¬¸ ìê¸°ì†Œê°œì„œ ì‘ì„± ì½”ì¹˜ì…ë‹ˆë‹¤.
        í†¤: {st.session_state.model_settings["tone"]}
        ìµœëŒ€ ê¸¸ì´: {st.session_state.model_settings["max_length"]}ì
        
        - êµ¬ì²´ì ì´ê³  ì‹¤ìš©ì ì¸ ì¡°ì–¸
        - ì˜ˆì‹œë¥¼ ë“¤ì–´ ì„¤ëª…
        - ì¹œê·¼í•˜ë©´ì„œë„ ì „ë¬¸ì ì¸ í†¤
        - ì´ëª¨ì§€ëŠ” ìµœì†Œí•œìœ¼ë¡œ ì‚¬ìš©"""
        
        if uploaded_file:
            try:
                if uploaded_file.name.endswith('.txt'):
                    content = uploaded_file.read().decode('utf-8')
                elif uploaded_file.name.endswith('.docx') and DOC_LIBS_AVAILABLE:
                    doc = Document(uploaded_file)
                    content = '\n'.join([p.text for p in doc.paragraphs])
                else:
                    content = "íŒŒì¼ì„ ì½ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤."
                
                user_input = f"ë‹¤ìŒ ìê¸°ì†Œê°œì„œë¥¼ ê²€í† í•˜ê³  ê°œì„ ì ì„ ì œì•ˆí•´ì£¼ì„¸ìš”:\n\n{content}\n\n{user_input}"
            except Exception as e:
                return f"íŒŒì¼ ì²˜ë¦¬ ì¤‘ ì˜¤ë¥˜: {e}"
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{input}")
        ])
        
        chain = LLMChain(llm=llm, prompt=prompt)
        response = chain.invoke({"input": user_input})
        
        return response.get("text", str(response))
        
    except Exception as e:
        return f"ì˜¤ë¥˜ê°€ ë°œìƒí–ˆìŠµë‹ˆë‹¤. ë‹¤ì‹œ ì‹œë„í•´ì£¼ì„¸ìš”.\n{str(e)}"

# ================= ëŒ€í™” ì €ì¥ =================
def save_conversation():
    content = ""
    for msg in st.session_state.messages:
        role = "ğŸ‘¤ ì‚¬ìš©ì" if msg["role"] == "user" else "ğŸ¤– AI ì½”ì¹˜"
        content += f"[{msg.get('time', '')}] {role}\n{msg['content']}\n\n"
    
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"ìì†Œì„œëŒ€í™”_{timestamp}"
    
    # ì„ íƒëœ í˜•ì‹ìœ¼ë¡œ ì €ì¥
    if st.session_state.save_format == "txt":
        file_data = content
        mime = "text/plain"
        ext = "txt"
    elif st.session_state.save_format == "docx" and DOC_LIBS_AVAILABLE:
        doc = Document()
        doc.add_heading('AI ìê¸°ì†Œê°œì„œ ì½”ì¹­ ëŒ€í™”', 0)
        for para in content.split('\n'):
            doc.add_paragraph(para)
        
        bio = io.BytesIO()
        doc.save(bio)
        file_data = bio.getvalue()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    else:
        file_data = content
        mime = "text/plain"
        ext = "txt"
    
    # ì €ì¥ ëª©ë¡ì— ì¶”ê°€
    st.session_state.saved_files.append({
        "name": f"{filename}.{ext}",
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "size": len(file_data),
        "data": file_data,
        "mime": mime
    })
    
    return f"{filename}.{ext}"

# ================= UI ë Œë”ë§ í•¨ìˆ˜ =================
def render_header():
    st.markdown(f'''
        <div class="chat-header">
            <div class="header-title">AI ìê¸°ì†Œê°œì„œ ì½”ì¹­</div>
        </div>
    ''', unsafe_allow_html=True)

def render_chat_tab():
    # ì±„íŒ… ë©”ì‹œì§€ í‘œì‹œ
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            st.markdown(f'''
                <div class="msg-row user">
                    <div class="msg-time">{msg.get("time", "")}</div>
                    <div class="msg-bubble user">{msg["content"]}</div>
                </div>
            ''', unsafe_allow_html=True)
        else:
            content_html = msg["content"].replace('\n', '<br>')
            st.markdown(f'''
                <div class="msg-row ai">
                    <div class="msg-bubble ai">{content_html}</div>
                    <div class="msg-time">{msg.get("time", "")}</div>
                </div>
            ''', unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # ì…ë ¥ ì˜ì—­ (ê³ ì •)
    with st.container():
        st.markdown('<div class="chat-input-section">', unsafe_allow_html=True)
        
        # ë¹ ë¥¸ ë‹µë³€
        st.markdown('<div class="quick-replies">', unsafe_allow_html=True)
        quick_replies = ["ê°€ì´ë“œë¼ì¸ ì•Œë ¤ì¤˜", "ìì†Œì„œ ì‹œì‘í•˜ê¸°", "ì²¨ì‚­ ë°›ê³  ì‹¶ì–´", "ì˜ˆì‹œ ë³´ì—¬ì¤˜"]
        cols = st.columns(len(quick_replies))
        for i, reply in enumerate(quick_replies):
            with cols[i]:
                if st.button(reply, key=f"quick_{i}"):
                    st.session_state.messages.append({
                        "role": "user",
                        "content": reply,
                        "time": datetime.datetime.now().strftime("%H:%M")
                    })
                    response = get_ai_response(reply)
                    st.session_state.messages.append({
                        "role": "ai",
                        "content": response,
                        "time": datetime.datetime.now().strftime("%H:%M")
                    })
                    st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
        
        # íŒŒì¼ ì—…ë¡œë“œ
        uploaded_file = st.file_uploader(
            "ğŸ“ íŒŒì¼ ì²¨ë¶€ (txt, docx)",
            type=['txt', 'docx'],
            label_visibility="visible",
            help="ìê¸°ì†Œê°œì„œ íŒŒì¼ì„ ì—…ë¡œë“œí•˜ì—¬ ì²¨ì‚­ë°›ì„ ìˆ˜ ìˆìŠµë‹ˆë‹¤."
        )
        
        # ë©”ì‹œì§€ ì…ë ¥
        with st.form("chat_form", clear_on_submit=True):
            col1, col2 = st.columns([5, 1])
            with col1:
                user_input = st.text_input(
                    "ë©”ì‹œì§€",
                    placeholder="ë©”ì‹œì§€ë¥¼ ì…ë ¥í•˜ì„¸ìš”...",
                    label_visibility="collapsed"
                )
            with col2:
                send = st.form_submit_button("ì „ì†¡")
            
            if send and user_input:
                # ì‚¬ìš©ì ë©”ì‹œì§€ ì¶”ê°€
                st.session_state.messages.append({
                    "role": "user",
                    "content": user_input,
                    "time": datetime.datetime.now().strftime("%H:%M")
                })
                
                # AI ì‘ë‹µ ìƒì„±
                with st.spinner("ì…ë ¥ ì¤‘..."):
                    response = get_ai_response(user_input, uploaded_file)
                
                st.session_state.messages.append({
                    "role": "ai",
                    "content": response,
                    "time": datetime.datetime.now().strftime("%H:%M")
                })
                
                st.rerun()
        
        st.markdown('</div>', unsafe_allow_html=True)

def render_settings_tab():
    st.markdown('<div class="settings-container">', unsafe_allow_html=True)
    
    st.title("âš™ï¸ ì„¤ì •")
    
    # API ì„¤ì •
    st.markdown('<div class="settings-section">', unsafe_allow_html=True)
    st.markdown('<div class="settings-title">ğŸ”‘ API ì„¤ì •</div>', unsafe_allow_html=True)
    
    api_key = st.text_input(
        "OpenAI API Key",
        value=st.session_state.api_key,
        type="password",
        placeholder="sk-...",
        help="OpenAI API í‚¤ë¥¼ ì…ë ¥í•˜ë©´ ë” ì •í™•í•œ AI ì‘ë‹µì„ ë°›ì„ ìˆ˜ ìˆìŠµë‹ˆë‹¤."
    )
    
    if api_key != st.session_state.api_key:
        st.session_state.api_key = api_key
        st.success("API í‚¤ê°€ ì €ì¥ë˜ì—ˆìŠµë‹ˆë‹¤!")
    
    st.info("ğŸ’¡ API í‚¤ê°€ ì—†ì–´ë„ ê¸°ë³¸ ê¸°ëŠ¥ì„ ì‚¬ìš©í•  ìˆ˜ ìˆìŠµë‹ˆë‹¤")
    st.markdown('</div>', unsafe_allow_html=True)
    
    # ëŒ€í™” ê´€ë¦¬
    st.markdown('<div class="settings-section">', unsafe_allow_html=True)
    st.markdown('<div class="settings-title">ğŸ’¬ ëŒ€í™” ê´€ë¦¬</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("ğŸ—‘ï¸ ëŒ€í™” ì´ˆê¸°í™”"):
            st.session_state.messages = [{
                "role": "ai",
                "content": "ì•ˆë…•í•˜ì„¸ìš”! AI ìê¸°ì†Œê°œì„œ ì½”ì¹˜ì…ë‹ˆë‹¤. ë¬´ì—‡ì„ ë„ì™€ë“œë¦´ê¹Œìš”?",
                "time": datetime.datetime.now().strftime("%H:%M")
            }]
            st.success("ëŒ€í™”ê°€ ì´ˆê¸°í™”ë˜ì—ˆìŠµë‹ˆë‹¤!")
            st.rerun()
    
    with col2:
        if st.button("ğŸ’¾ ëŒ€í™” ì €ì¥"):
            filename = save_conversation()
            st.success(f"{filename} ì €ì¥ë¨!")
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

def render_advanced_settings_tab():
    st.markdown('<div class="settings-container">', unsafe_allow_html=True)
    
    st.title("ğŸ”§ ì„¸ë¶€ì„¤ì •")
    
    # AI ëª¨ë¸ ì„¤ì •
    st.markdown('<div class="settings-section">', unsafe_allow_html=True)
    st.markdown('<div class="settings-title">ğŸ¤– AI ëª¨ë¸ ì„¤ì •</div>', unsafe_allow_html=True)
    
    st.session_state.model_settings["temperature"] = st.slider(
        "ì°½ì˜ì„± (Temperature)",
        min_value=0.0,
        max_value=1.0,
        value=st.session_state.model_settings["temperature"],
        step=0.1,
        help="ê°’ì´ ë†’ì„ìˆ˜ë¡ ë” ì°½ì˜ì ì¸ ë‹µë³€ì„ ìƒì„±í•©ë‹ˆë‹¤."
    )
    
    st.session_state.model_settings["max_length"] = st.number_input(
        "ìµœëŒ€ ì‘ë‹µ ê¸¸ì´ (ì)",
        min_value=100,
        max_value=3000,
        value=st.session_state.model_settings["max_length"],
        step=100,
        help="AI ì‘ë‹µì˜ ìµœëŒ€ ê¸¸ì´ë¥¼ ì„¤ì •í•©ë‹ˆë‹¤."
    )
    
    st.session_state.model_settings["tone"] = st.selectbox(
        "ì‘ë‹µ í†¤",
        ["professional", "friendly", "casual", "formal"],
        index=["professional", "friendly", "casual", "formal"].index(st.session_state.model_settings["tone"]),
        help="AIì˜ ì‘ë‹µ ìŠ¤íƒ€ì¼ì„ ì„ íƒí•©ë‹ˆë‹¤."
    )
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # ì €ì¥ ì„¤ì •
    st.markdown('<div class="settings-section">', unsafe_allow_html=True)
    st.markdown('<div class="settings-title">ğŸ’¾ ì €ì¥ ì„¤ì •</div>', unsafe_allow_html=True)
    
    st.session_state.save_format = st.selectbox(
        "ê¸°ë³¸ ì €ì¥ í˜•ì‹",
        ["txt", "docx", "pdf"],
        index=["txt", "docx", "pdf"].index(st.session_state.save_format),
        help="ëŒ€í™” ì €ì¥ ì‹œ ì‚¬ìš©í•  íŒŒì¼ í˜•ì‹ì„ ì„ íƒí•©ë‹ˆë‹¤."
    )
    
    st.info("ğŸ“Œ ì €ì¥ëœ íŒŒì¼ì€ 'ì €ì¥ì†Œ' íƒ­ì—ì„œ í™•ì¸í•  ìˆ˜ ìˆìŠµë‹ˆë‹¤")
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

def render_storage_tab():
    st.markdown('<div class="storage-container">', unsafe_allow_html=True)
    
    st.title("ğŸ“ ì €ì¥ì†Œ")
    
    if not st.session_state.saved_files:
        st.info("ì €ì¥ëœ íŒŒì¼ì´ ì—†ìŠµë‹ˆë‹¤. ëŒ€í™”ë¥¼ ì €ì¥í•˜ë ¤ë©´ ì„¤ì • íƒ­ì„ ì´ìš©í•˜ì„¸ìš”.")
    else:
        st.write(f"ì´ {len(st.session_state.saved_files)}ê°œì˜ íŒŒì¼ì´ ì €ì¥ë˜ì–´ ìˆìŠµë‹ˆë‹¤.")
        
        for i, file in enumerate(st.session_state.saved_files):
            with st.container():
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    st.markdown(f'''
                        <div class="file-item">
                            <div class="file-info">
                                <div class="file-name">ğŸ“„ {file["name"]}</div>
                                <div class="file-date">{file["date"]} Â· {file["size"]} bytes</div>
                            </div>
                        </div>
                    ''', unsafe_allow_html=True)
                
                with col2:
                    st.download_button(
                        label="ë‹¤ìš´ë¡œë“œ",
                        data=file["data"],
                        file_name=file["name"],
                        mime=file["mime"],
                        key=f"download_{i}_{file['name']}"
                    )
    
    # ì¼ê´„ ì‚­ì œ
    if st.session_state.saved_files:
        st.markdown("---")
        if st.button("ğŸ—‘ï¸ ëª¨ë“  íŒŒì¼ ì‚­ì œ"):
            st.session_state.saved_files = []
            st.success("ëª¨ë“  íŒŒì¼ì´ ì‚­ì œë˜ì—ˆìŠµë‹ˆë‹¤!")
            st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

# ================= ë©”ì¸ ì•± =================
def main():
    # í—¤ë”
    render_header()
    
    # íƒ­ ìƒì„± (Streamlit ë‚´ì¥ íƒ­ ì‚¬ìš©)
    tab1, tab2, tab3, tab4 = st.tabs(["ğŸ’¬ ëŒ€í™”", "âš™ï¸ ì„¤ì •", "ğŸ”§ ì„¸ë¶€ì„¤ì •", "ğŸ“ ì €ì¥ì†Œ"])
    
    with tab1:
        render_chat_tab()
    
    with tab2:
        render_settings_tab()
    
    with tab3:
        render_advanced_settings_tab()
    
    with tab4:
        render_storage_tab()

# í”„ë¡œê·¸ë¨ ì§„ì…ì 
if __name__ == "__main__":
    main()
