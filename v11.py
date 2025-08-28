# app.py
# =========================================================
# AI 자기소개서 코칭 - Streamlit UI (v11)
# =========================================================
# 설치: pip install streamlit python-docx reportlab langchain langchain-openai python-dotenv
# 실행: streamlit run v11.py
# =========================================================

import os, io, datetime, json
from typing import Optional, List, Dict
import streamlit as st

# ===== 문서 생성 라이브러리 (선택) =====
try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    DOC_LIBS_AVAILABLE = True
except:
    DOC_LIBS_AVAILABLE = False

# ===== LangChain (선택) =====
try:
    from langchain_openai import ChatOpenAI
    from langchain.prompts import ChatPromptTemplate
    from langchain.chains import LLMChain
    LANGCHAIN_AVAILABLE = True
except:
    LANGCHAIN_AVAILABLE = False

# ================= 페이지 설정 및 기본 스타일 =================
st.set_page_config(
    page_title="AI 자기소개서 코칭",
    page_icon="💬",
    layout="wide",
    initial_sidebar_state="collapsed"
)

MAIN_COLOR = "#22C55E"       # 메인 초록색
SUB_COLOR = "#DCFCE7"        # 사용자 말풍선 배경
BOT_COLOR = "#F3F4F6"        # 챗봇 말풍선 배경
BG_COLOR = "#F5FBFB"         # 전체 배경색

st.markdown(
    f"""
    <style>
        body {{
            background-color: {BG_COLOR};
        }}

        .chat-header-title {{
            color: white;
            font-weight: 600;
        }}
        .bottom-nav {{
            position: fixed;
            left: 0;
            right: 0;
            bottom: 0;
            background: white;
            border-top: 1px solid #e0e0e0;
            padding: 4px 8px;
        }}
        .bottom-nav button {{
            width: 100%;
            background: transparent;
            border: none;
            color: {MAIN_COLOR};
            font-size: 14px;
        }}
        .bottom-nav .active {{
            color: white;
            background: {MAIN_COLOR};
            border-radius: 12px;
        }}
        .nav-icon {{
            font-size: 20px;
            display: block;
        }}
        .onboard-wrapper {{
            text-align: center;
            padding: 60px 20px;
        }}
        .onboard-circle {{
            width: 120px;
            height: 120px;
            border-radius: 60px;
            background: {SUB_COLOR};
            margin: 0 auto 24px auto;
            display:flex;
            align-items:center;
            justify-content:center;
            font-size:32px;
        }}

        .stMainBlockContainer {{
            // padding: 0;
        }}

        .stVerticalBlock {{
            // gap: 0;
        }}

        .stAppHeader {{
            display: none;
        }}

        .header {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
        }}
    </style>
    """,
    unsafe_allow_html=True,
)

# ================= 세션 초기화 =================
if "messages" not in st.session_state:
    st.session_state.messages = []
    st.session_state.messages.append({
        "role": "ai",
        "content": "안녕하세요! AI 자기소개서 코치입니다. 무엇을 도와드릴까요?",
        "time": datetime.datetime.now().strftime("%H:%M")
    })

if "current_tab" not in st.session_state:
    st.session_state.current_tab = "대화"

if "api_key" not in st.session_state:
    st.session_state.api_key = os.getenv("OPENAI_API_KEY", "")

if "saved_files" not in st.session_state:
    st.session_state.saved_files = []

if "basic_settings" not in st.session_state:
    st.session_state.basic_settings = {
        "model": "GPT-4 (무료)",
        "tone": "전문적",
        "length": 800,
    }

if "advanced_settings" not in st.session_state:
    st.session_state.advanced_settings = {
        "creativity": 0.5,
        "polish": 0.5,
        "auto_save": True,
        "smart_edit": True,
        "export_format": "PDF 문서",
    }

if "show_saved" not in st.session_state:
    st.session_state.show_saved = False

# ================= 가이드라인 응답 =================
def get_guideline() -> str:
    return """📝 **AI 자기소개서 입력 가이드라인**

**1. 구체적으로 질문하기**
✅ "마케팅 직무 신입 자기소개서 도입부 작성해줘"
❌ "자소서 써줘"

**2. 배경 정보 제공하기**
• 지원 회사와 직무
• 본인의 주요 경험
• 강조하고 싶은 역량

**3. 효과적인 질문 예시**
• "고객 서비스 경험을 영업직무에 연결하는 방법"
• "프로젝트 경험을 STAR 기법으로 정리해줘"
• "IT 기업 지원동기 작성 도와줘"

**4. 첨삭 요청 방법**
• 작성한 문장을 복사 후 "이 내용 첨삭해줘"
• 파일 업로드 후 "구체성 높여줘"
• "이 문장 더 임팩트 있게 수정해줘"

**5. 단계별 접근**
1️⃣ 전체 구조 잡기
2️⃣ 각 문단 작성
3️⃣ 표현 다듬기
4️⃣ 최종 검토

💡 **Tip**: 한 번에 모든 걸 해결하려 하지 말고, 단계별로 질문하세요!"""

# ================= AI 응답 생성 =================
def get_ai_response(user_input: str, uploaded_file=None) -> str:
    guideline_keywords = ["가이드", "가이드라인", "도움말", "사용법", "어떻게"]
    if any(keyword in user_input for keyword in guideline_keywords):
        return get_guideline()

    if not st.session_state.api_key or not LANGCHAIN_AVAILABLE:
        templates = {
            "default": """자기소개서 작성을 도와드리겠습니다!

구체적으로 알려주시면 더 정확한 도움을 드릴 수 있어요:
• 어떤 직무에 지원하시나요?
• 어떤 부분이 어려우신가요?
• 특별히 강조하고 싶은 경험이 있나요?""",
            "첨삭": """자기소개서 첨삭 포인트를 알려드릴게요:

✅ 구체적인 숫자와 성과 포함
✅ 직무와 연관된 경험 강조
✅ 문장은 간결하고 명확하게
✅ 진정성 있는 지원동기

파일을 업로드하거나 내용을 보내주시면 더 자세히 봐드릴게요!""",
            "시작": """자기소개서 작성을 시작해볼까요?

**Step 1. 기본 정보**
• 지원 회사:
• 지원 직무:
• 경력 구분: (신입/경력)

이 정보를 알려주시면 맞춤형으로 도와드릴게요!""",
            "예시": """다음은 간단한 자기소개서 예시입니다:

"문제 해결 능력을 바탕으로 한 프로젝트 경험을 통해 팀에 기여했던 사례가 있습니다."

이와 같은 방식으로 경험을 구체적으로 설명해보세요!""",
        }
        if "첨삭" in user_input or "수정" in user_input:
            return templates["첨삭"]
        elif "시작" in user_input or "처음" in user_input:
            return templates["시작"]
        elif "예시" in user_input:
            return templates["예시"]
        else:
            return templates["default"]

    try:
        model_map = {
            "GPT-4 (무료)": "gpt-4o-mini",
            "GPT-4": "gpt-4o",
            "GPT-3.5": "gpt-3.5-turbo",
        }
        selected_model = st.session_state.basic_settings.get("model", "GPT-4 (무료)")
        model_name = model_map.get(selected_model, "gpt-4o-mini")
        llm = ChatOpenAI(
            api_key=st.session_state.api_key,
            model=model_name,
            temperature=st.session_state.advanced_settings["creativity"],
        )

        system_prompt = f"""당신은 전문 자기소개서 작성 코치입니다.
        톤: {st.session_state.basic_settings['tone']}
        최대 길이: {st.session_state.basic_settings['length']}자

        - 구체적이고 실용적인 조언
        - 예시를 들어 설명
        - 친근하면서도 전문적인 톤
        - 이모지는 최소한으로 사용"""

        if uploaded_file:
            try:
                if uploaded_file.name.endswith('.txt'):
                    content = uploaded_file.read().decode('utf-8')
                elif uploaded_file.name.endswith('.docx') and DOC_LIBS_AVAILABLE:
                    doc = Document(uploaded_file)
                    content = '\n'.join([p.text for p in doc.paragraphs])
                else:
                    content = "파일을 읽을 수 없습니다."
                user_input = f"다음 자기소개서를 검토하고 개선점을 제안해주세요:\n\n{content}\n\n{user_input}"
            except Exception as e:
                return f"파일 처리 중 오류: {e}"

        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{input}")
        ])
        chain = LLMChain(llm=llm, prompt=prompt)
        response = chain.invoke({"input": user_input})
        return response.get("text", str(response))
    except Exception as e:
        return f"오류가 발생했습니다. 다시 시도해주세요.\n{str(e)}"

# ================= 대화 저장 =================
def save_conversation():
    content = ""
    for msg in st.session_state.messages:
        role = "👤 사용자" if msg["role"] == "user" else "🤖 AI 코치"
        content += f"[{msg.get('time', '')}] {role}\n{msg['content']}\n\n"

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"자소서대화_{timestamp}"
    export = st.session_state.advanced_settings.get("export_format", "텍스트 파일")

    if export == "PDF 문서" and DOC_LIBS_AVAILABLE:
        bio = io.BytesIO()
        doc = SimpleDocTemplate(bio, pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph(p, styles["Normal"]) for p in content.split('\n')]
        doc.build(story)
        file_data = bio.getvalue()
        mime = "application/pdf"
        ext = "pdf"
    elif export == "Word 문서" and DOC_LIBS_AVAILABLE:
        doc = Document()
        doc.add_heading('AI 자기소개서 코칭 대화', 0)
        for para in content.split('\n'):
            doc.add_paragraph(para)
        bio = io.BytesIO()
        doc.save(bio)
        file_data = bio.getvalue()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    elif export == "HTML 문서":
        file_data = f"<html><body><pre>{content}</pre></body></html>"
        mime = "text/html"
        ext = "html"
    else:
        file_data = content
        mime = "text/plain"
        ext = "txt"

    st.session_state.saved_files.append({
        "name": f"{filename}.{ext}",
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "size": len(file_data),
        "data": file_data,
        "mime": mime
    })

    return f"{filename}.{ext}"

##########################################
# UI 렌더링 함수
##########################################

def render_header(title: str) -> None:
    st.markdown(
        f"<div class='header' style='background:{MAIN_COLOR}; padding:12px; text-align:center; color:white; font-weight:600'>{title}</div>",
        unsafe_allow_html=True,
    )


def render_bottom_nav() -> None:
    st.markdown("<div class='bottom-nav'>", unsafe_allow_html=True)
    cols = st.columns(4)
    tabs = ["대화", "설정", "세부 설정", "계정"]
    icons = ["💬", "⚙️", "🛠️", "👤"]
    for col, tab, icon in zip(cols, tabs, icons):
        label = f"{icon} {tab}"
        if col.button(label, key=f"nav_{tab}", use_container_width=True):
            st.session_state.current_tab = tab
            st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)


def render_onboarding():
    render_header("AI 자기소개서")
    st.markdown(
        "<div class='onboard-wrapper'>"\
        "<div class='onboard-circle'>✍️</div>"\
        "<h3>AI 자기소개서</h3>"\
        "<p>AI와 대화하면서 나만의 탄탄한 자기소개서를 완성하세요.</p>"\
        "<ol style='text-align:left; display:inline-block;'>"\
        "<li>AI와 대화를 통해 작성의 방향을 잡아</li>"\
        "<li>궁금한 질문은 언제든지 톡! 작성 톤을 설정하고</li>"\
        "<li>완벽하게 마무리된 자기소개서를 완성</li>"\
        "</ol>"\
        "</div>",
        unsafe_allow_html=True,
    )
    if st.button("시작하기", use_container_width=True):
        st.session_state.started = True
        st.session_state.current_tab = "대화"
        st.rerun()


def render_chat_tab():
    render_header("AI 대화")
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            st.markdown(
                f"<div style='text-align:right; background:{SUB_COLOR}; padding:10px; border-radius:18px; margin:4px 0'>{msg['content']}</div>",
                unsafe_allow_html=True,
            )
        else:
            content_html = msg["content"].replace("\n", "<br>")
            st.markdown(
                f"<div style='text-align:left; background:{BOT_COLOR}; padding:10px; border-radius:18px; margin:4px 0'>{content_html}</div>",
                unsafe_allow_html=True,
            )
    st.write("---")
    uploaded_file = st.file_uploader("📎 파일 첨부 (txt, docx)", type=["txt", "docx"])
    col1, col2, col3, col4 = st.columns([5, 1, 1, 1])
    with col1:
        user_input = st.text_input("메시지", placeholder="메시지를 입력하세요...", label_visibility="collapsed")
    with col2:
        send = st.button("전송")
    with col3:
        save = st.button("저장하기")
    with col4:
        if st.button("📂"):
            st.session_state.show_saved = not st.session_state.get("show_saved", False)
    if send and user_input:
        st.session_state.messages.append({
            "role": "user",
            "content": user_input,
            "time": datetime.datetime.now().strftime("%H:%M"),
        })
        with st.spinner("답변 생성 중..."):
            response = get_ai_response(user_input, uploaded_file)
        st.session_state.messages.append({
            "role": "ai",
            "content": response,
            "time": datetime.datetime.now().strftime("%H:%M"),
        })
        st.rerun()
    if save:
        filename = save_conversation()
        st.success(f"{filename} 저장됨!")
    if st.session_state.get("show_saved", False):
        st.markdown("---")
        if not st.session_state.saved_files:
            st.info("저장된 파일이 없습니다.")
        else:
            for i, file in enumerate(st.session_state.saved_files):
                st.write(f"📄 {file['name']} ({file['date']}, {file['size']} bytes)")
                st.download_button(
                    label="다운로드",
                    data=file["data"],
                    file_name=file["name"],
                    mime=file["mime"],
                    key=f"download_{i}_{file['name']}",
                )
            if st.button("🗑️ 모든 파일 삭제"):
                st.session_state.saved_files = []
                st.success("모든 파일이 삭제되었습니다!")
                st.session_state.show_saved = False
    render_bottom_nav()


def render_settings_tab():
    render_header("기본 설정")
    models = ["GPT-4 (무료)", "GPT-4", "GPT-3.5"]
    st.session_state.basic_settings["model"] = st.selectbox(
        "AI 모델 선택",
        models,
        index=models.index(st.session_state.basic_settings.get("model", models[0])),
    )
    tones = ["전문적", "친근한", "격식 있는", "캐주얼"]
    st.session_state.basic_settings["tone"] = st.selectbox(
        "작성 톤",
        tones,
        index=tones.index(st.session_state.basic_settings.get("tone", tones[0])),
    )
    st.session_state.basic_settings["length"] = st.slider(
        "글자 수",
        min_value=300,
        max_value=2000,
        value=st.session_state.basic_settings.get("length", 800),
    )
    render_bottom_nav()


def render_advanced_settings_tab():
    render_header("세부 설정")
    st.session_state.advanced_settings["creativity"] = st.slider(
        "창의성",
        0.0,
        1.0,
        value=st.session_state.advanced_settings.get("creativity", 0.5),
    )
    st.session_state.advanced_settings["polish"] = st.slider(
        "완성 수준",
        0.0,
        1.0,
        value=st.session_state.advanced_settings.get("polish", 0.5),
    )
    st.markdown("---")
    st.session_state.advanced_settings["auto_save"] = st.toggle(
        "자동 저장", value=st.session_state.advanced_settings.get("auto_save", True)
    )
    st.session_state.advanced_settings["smart_edit"] = st.toggle(
        "스마트 편집", value=st.session_state.advanced_settings.get("smart_edit", True)
    )
    st.markdown("---")
    export_options = ["PDF 문서", "Word 문서", "텍스트 파일", "HTML 문서"]
    st.session_state.advanced_settings["export_format"] = st.selectbox(
        "내보내기 설정",
        export_options,
        index=export_options.index(st.session_state.advanced_settings.get("export_format", "PDF 문서")),
    )
    render_bottom_nav()


def render_account_tab():
    render_header("계정")
    key = st.text_input(
        "OpenAI API Key",
        value=st.session_state.api_key,
        type="password",
        placeholder="sk-...",
    )
    if key != st.session_state.api_key:
        st.session_state.api_key = key
        st.success("API 키가 저장되었습니다!")
    render_bottom_nav()

# ================= 메인 앱 =================
def main():
    if "started" not in st.session_state:
        st.session_state.started = False
    if not st.session_state.started:
        render_onboarding()
        return
    page = st.session_state.get("current_tab", "대화")
    if page == "대화":
        render_chat_tab()
    elif page == "설정":
        render_settings_tab()
    elif page == "세부 설정":
        render_advanced_settings_tab()
    else:
        render_account_tab()

if __name__ == "__main__":
    main()
